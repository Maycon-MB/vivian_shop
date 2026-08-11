# -*- coding: utf-8 -*-
"""
Gera o documento de acompanhamento que a cliente e o Maycon usam juntos.

    python scripts/gerar_documento_cliente.py

Produz docs/combinado-do-projeto.docx, para subir no Google Drive e
compartilhar com a cliente com permissão de edição.

Por que um documento editável e não a página do site: a página é só
leitura — GitHub Pages é estático e não guarda o que alguém escreve. A
cliente precisa digitar as respostas dela e anotar o que for lembrando.

Por que Google Docs e não planilha: ela vive no WhatsApp. Documento se lê
como texto, abre no celular, e digitar embaixo de uma pergunta é óbvio.
Planilha exige entender linha, coluna e onde pode escrever.

Este script gera o documento inicial. Depois que a cliente começar a
escrever nele, o Google Docs passa a ser a fonte — regerar apagaria as
respostas dela. A partir daí, o script serve de referência do formato.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys, os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Paleta da identidade — ver docs/superpowers/specs/2026-08-10-identidade-visual-design.md
INK = RGBColor(0x12, 0x30, 0x5B)
CHALK = RGBColor(0x2E, 0x9B, 0x96)
SOFT = RGBColor(0x6B, 0x7C, 0x8F)
HEART = RGBColor(0xC4, 0x43, 0x6B)

FUNDO_PERGUNTA = "FFF8D6"   # amarelo bem claro, do marca-texto
FUNDO_PRONTO = "E8F5F1"     # verde-água bem claro
FUNDO_NEUTRO = "F4F7FB"


def sombrear(paragrafo, cor_hex):
    """Pinta o fundo de um parágrafo inteiro."""
    pPr = paragrafo._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), cor_hex)
    pPr.append(shd)


def texto(doc, conteudo, tamanho=11, cor=None, negrito=False, italico=False,
          espaco_antes=0, espaco_depois=6, fundo=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(espaco_antes)
    p.paragraph_format.space_after = Pt(espaco_depois)
    run = p.add_run(conteudo)
    run.font.size = Pt(tamanho)
    run.font.bold = negrito
    run.font.italic = italico
    run.font.color.rgb = cor if cor else INK
    if fundo:
        sombrear(p, fundo)
    return p


def titulo(doc, conteudo, nivel=1):
    tamanhos = {0: 24, 1: 16, 2: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20 if nivel else 0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(conteudo)
    run.font.size = Pt(tamanhos[nivel])
    run.font.bold = True
    run.font.color.rgb = INK
    return p


def campo_resposta(doc, rotulo="Sua resposta:"):
    """O espaço onde ela digita. Fundo amarelo claro para não ter dúvida."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(rotulo + " ")
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = INK
    r2 = p.add_run("(escreva aqui)")
    r2.font.size = Pt(11)
    r2.font.italic = True
    r2.font.color.rgb = SOFT
    sombrear(p, FUNDO_PERGUNTA)
    return p


def linha_fina(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("─" * 46)
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xA8, 0xC6, 0xE8)


# ─────────────────────────────────────────────────────────────────────────
ANDAMENTO = [
    ("Pronto", FUNDO_PRONTO, [
        ("Identidade visual", "Cores e letras da loja, tiradas das suas duas logos do Elo7."),
        ("Telas da loja para você ver", "Catálogo, carrinho e regras de venda, com produtos de exemplo."),
        ("Explicação de como tudo funciona", "O passo a passo de uma venda, do pagamento até o envio."),
        ("Este documento", "Para nós dois acompanharmos o que falta."),
    ]),
    ("Fazendo agora", FUNDO_PERGUNTA, [
        ("Loja de verdade, com os seus produtos", "Sair do exemplo e montar o catálogo real, com as suas fotos e preços."),
    ]),
    ("Vem a seguir", FUNDO_NEUTRO, [
        ("Pagamento por Pix e cartão", "Ligar a loja à sua conta do Mercado Pago, para o dinheiro cair direto para você."),
        ("Cálculo de frete e etiqueta", "Correios e Jadlog calculando na hora, com etiqueta e declaração saindo prontas."),
        ("Entrega do material pedagógico", "O arquivo indo sozinho para o e-mail assim que o pagamento é aprovado."),
        ("Seu painel de verdade", "Cadastrar produto, ver pedidos e gerar etiqueta, tudo pelo celular."),
    ]),
    ("Mais para frente", FUNDO_NEUTRO, [
        ("Endereço próprio e loja no ar", "Registrar o endereço do site e abrir a loja para o público."),
        ("Posts no Instagram", "Começa depois que a loja estiver vendendo, para o post ter para onde mandar as pessoas."),
        ("Anúncios pagos", "Por último, quando já soubermos o que mais vende."),
    ]),
]

PERGUNTAS = [
    ("O visual está parecendo a sua marca?",
     "Tirei as cores das suas duas logos: o verde-água e o rosa vieram do \"Feito para você\", o amarelo veio do \"Projeto Educar\", e o azul que as duas tinham virou a cor do texto.",
     "Essa é a hora de mudar. Depois que a loja estiver construída em cima disso, fica bem mais caro."),

    ("Como são as atividades pedagógicas?",
     "Preciso entender o formato para montar a entrega automática.",
     "São PDF para a pessoa imprimir em casa? Cada compra é um arquivo só ou um pacote com vários? E o preço é por arquivo?"),

    ("Posso colocar o nome de quem comprou no arquivo?",
     "Material digital corre o risco de ser repassado em grupo de WhatsApp. O jeito mais usado de evitar é o arquivo sair com o nome de quem comprou escrito pequeno em cada página.",
     "Não atrapalha quem pagou, mas desanima quem ia repassar. Você topa? É melhor decidir agora, porque depois eu teria que refazer tudo que já foi vendido."),

    ("Quantas vendas por mês, mais ou menos?",
     "Pode ser um chute, com base nos últimos meses no Elo7.",
     "Serve para eu deixar a loja do tamanho certo e o custo mensal baixo."),

    ("Tudo bem o seu endereço aparecer na etiqueta?",
     "Os Correios exigem remetente real, então o seu endereço em Vila Valqueire vai impresso em toda etiqueta que sair. Quem compra consegue ver.",
     "Era assim no Elo7 também, mas prefiro que você saiba. Se quiser usar outro endereço, dá para trocar."),

    ("Falta alguma coisa no painel desde o primeiro dia?",
     "No começo o painel vai cadastrar produto, mostrar pedidos e gerar etiqueta. Chat dentro do painel, marketing automático e gráficos de venda ficam para depois.",
     "Prefiro entregar a loja vendendo cedo e ir crescendo com ela. Mas se alguma dessas coisas faz falta desde já para você, me fala."),
]

COMBINADO = [
    ("Nome da loja",
     "\"Feito para você! Personalizados\", com as duas linhas dentro, em vez de dois sites separados.", None),
    ("As duas linhas",
     "Papelaria personalizada, feita sob encomenda, e papelaria pedagógica, digital. Cada uma tem sua cor na loja: verde-água e amarelo.", None),
    ("Pedido mínimo",
     "Mínimo de 10 unidades de cada produto. Não dá para comprar 1 caneca — o mínimo são 10 canecas. Quem quiser dois modelos leva 10 de cada.", None),
    ("Prazo de produção",
     "5 dias úteis, contados de quando o pagamento é confirmado. O prazo aparece no produto, antes de a pessoa comprar.", None),
    ("Compras separadas",
     "Material digital e personalizado não vão na mesma compra. Quem quiser os dois faz dois pedidos.",
     "A declaração de conteúdo precisa bater com o que está dentro da caixa, e um arquivo digital declarado seria um item que não está na embalagem."),
    ("Transportadoras",
     "Correios e Jadlog, as mesmas que você já usava. Aparecem lado a lado com preço e prazo, e quem compra escolhe.", None),
    ("Declaração de conteúdo",
     "Continua existindo. Sai junto com a etiqueta, já preenchida com os itens, a quantidade e o valor. Você imprime, assina e leva.",
     "Pelos Correios está certo. Pela Jadlog ainda estou confirmando se aceitam a declaração ou se pedem nota fiscal."),
]

DECIDI_POR_CONTA = [
    ("Medidas para calcular o frete",
     "Você cadastra o peso e o tamanho do pacote fechado de 10, e não da peça solta.",
     "Como o mínimo é 10 unidades, é isso que você realmente despacha. Pesa uma vez, numa balança de cozinha, e nunca mais pensa nisso."),
    ("Como o material digital chega",
     "Por um link no e-mail, que vale por 7 dias. Não vai anexado.",
     "Arquivo anexado costuma cair na caixa de spam e tem limite de tamanho."),
    ("Aviso por WhatsApp",
     "No começo, o e-mail sai sozinho e o WhatsApp é um botão que abre a conversa com a mensagem já escrita, para você só enviar.",
     "Para ser totalmente automático, seria preciso um número dedicado que sai do WhatsApp normal. Você perderia o número que usa para atender pessoalmente — e esse atendimento direto é uma das melhores coisas que a sua loja tem."),
    ("Como funciona o pagamento",
     "A pessoa paga sem sair da sua loja, em vez de ser mandada para a tela do Mercado Pago.",
     "Menos gente desiste no meio do caminho. O dinheiro cai igual na sua conta."),
    ("Custo para manter a loja no ar",
     "Hospedagem, banco de dados e envio de e-mail ficam nos planos gratuitos. O único custo fixo é o endereço do site, cerca de R$ 40 por ano.",
     "Cabe com folga no valor mensal que combinamos. Se um dia o movimento crescer a ponto de sair do gratuito, eu te aviso antes de qualquer conta chegar."),
]

# ─────────────────────────────────────────────────────────────────────────
doc = Document()

for secao in doc.sections:
    secao.top_margin = Cm(2.2)
    secao.bottom_margin = Cm(2.2)
    secao.left_margin = Cm(2.4)
    secao.right_margin = Cm(2.4)

estilo = doc.styles["Normal"]
estilo.font.name = "Calibri"
estilo.font.size = Pt(11)
estilo.font.color.rgb = INK

# ── Capa ──
titulo(doc, "Feito para você! Personalizados", nivel=0)
texto(doc, "Acompanhamento do projeto da loja online", tamanho=12, cor=SOFT, espaco_depois=18)

texto(doc, "Para que serve este documento", tamanho=12, negrito=True, espaco_depois=6)
texto(doc,
      "Conversa de WhatsApp se perde. Aqui fica tudo num lugar só: o que já está pronto, "
      "o que estou fazendo agora, e o que ainda depende de você.",
      espaco_depois=6)
texto(doc,
      "Você pode escrever direto neste documento. Onde tiver um espaço amarelo, é para você "
      "responder. Não precisa responder tudo de uma vez, nem na ordem.",
      espaco_depois=6)
texto(doc,
      "E se lembrar de alguma coisa que quer na loja, escreve no final, na parte de anotações. "
      "Assim nenhum de nós dois esquece.",
      espaco_depois=6)

texto(doc, "Última atualização: 11 de agosto de 2026", tamanho=10, cor=SOFT, italico=True, espaco_antes=10)
texto(doc, "Para ver as telas: maycon-mb.github.io/vivian_shop", tamanho=10, cor=SOFT, espaco_depois=4)

doc.add_page_break()

# ── Andamento ──
titulo(doc, "Como está a construção", nivel=1)
texto(doc,
      "O que você vê nas telas hoje ainda é uma amostra, com produtos e valores de exemplo. "
      "A loja que vende de verdade é o que está sendo construído.",
      cor=SOFT, espaco_depois=14)

for fase, cor_fundo, itens in ANDAMENTO:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("  " + fase.upper() + "  ")
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = INK
    sombrear(p, cor_fundo)

    for nome, detalhe in itens:
        pi = doc.add_paragraph(style="List Bullet")
        pi.paragraph_format.space_after = Pt(3)
        rn = pi.add_run(nome + " — ")
        rn.font.bold = True
        rn.font.size = Pt(11)
        rn.font.color.rgb = INK
        rd = pi.add_run(detalhe)
        rd.font.size = Pt(11)
        rd.font.color.rgb = SOFT

doc.add_page_break()

# ── Perguntas ──
titulo(doc, "O que eu preciso saber", nivel=1)
texto(doc,
      "São seis perguntas. Responda no espaço amarelo embaixo de cada uma, aos poucos, "
      "na ordem que quiser.",
      cor=SOFT, espaco_depois=16)

for i, (pergunta, contexto, remate) in enumerate(PERGUNTAS, start=1):
    titulo(doc, f"{i}. {pergunta}", nivel=2)
    texto(doc, contexto, espaco_depois=4)
    if remate:
        texto(doc, remate, cor=SOFT, espaco_depois=6)
    campo_resposta(doc)

doc.add_page_break()

# ── Combinado ──
titulo(doc, "O que já está combinado", nivel=1)
texto(doc,
      "Isso aqui você já confirmou e está valendo. Se alguma coisa estiver diferente do que "
      "você quis dizer, escreve embaixo que eu corrijo.",
      cor=SOFT, espaco_depois=14)

for assunto, decisao, porque in COMBINADO:
    titulo(doc, assunto, nivel=2)
    texto(doc, decisao, espaco_depois=3)
    if porque:
        texto(doc, "Por quê: " + porque, tamanho=10, cor=SOFT, espaco_depois=4)
    linha_fina(doc)

campo_resposta(doc, "Alguma coisa aqui ficou diferente do que você quis dizer?")

doc.add_page_break()

# ── Decidi por conta ──
titulo(doc, "O que eu decidi sem te perguntar", nivel=1)
texto(doc,
      "Coisas técnicas que não precisam te tomar tempo. Estão aqui só para você saber — "
      "se discordar de alguma, é só falar.",
      cor=SOFT, espaco_depois=14)

for assunto, decisao, porque in DECIDI_POR_CONTA:
    titulo(doc, assunto, nivel=2)
    texto(doc, decisao, espaco_depois=3)
    if porque:
        texto(doc, "Por quê: " + porque, tamanho=10, cor=SOFT, espaco_depois=4)
    linha_fina(doc)

campo_resposta(doc, "Discorda de alguma?")

doc.add_page_break()

# ── Anotações ──
titulo(doc, "Suas anotações", nivel=1)
texto(doc,
      "Espaço livre. Se lembrar de alguma coisa que quer na loja, de um produto, de um jeito "
      "de fazer, de algo que te incomodava no Elo7 — escreve aqui. Não precisa estar organizado, "
      "eu leio e a gente conversa.",
      cor=SOFT, espaco_depois=12)

for _ in range(14):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(" ")
    r.font.size = Pt(11)
    sombrear(p, FUNDO_PERGUNTA)

os.makedirs("docs", exist_ok=True)
doc.save("docs/combinado-do-projeto.docx")
print("docs/combinado-do-projeto.docx")
