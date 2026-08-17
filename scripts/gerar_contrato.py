# -*- coding: utf-8 -*-
"""
Transforma o contrato em um .docx para mandar pela Vivian.

    python scripts/gerar_contrato.py

Existe porque markdown é ótimo para eu manter e péssimo para ela ler. Ela
vai abrir no celular, pelo WhatsApp, e precisa de um documento que pareça
um contrato — com título, cláusulas numeradas e linha de assinatura.

O arquivo de origem fica fora deste repositório, junto com os dados
pessoais dela. O que está versionado aqui é só o modelo em branco.
"""

import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor

PASTA = os.path.join(os.path.expanduser('~'), 'Documents', 'vivian-contrato')
ORIGEM = os.path.join(PASTA, 'contrato-vivian.md')
DESTINO = os.path.join(PASTA, 'Contrato - Loja Feito para Voce.docx')

TINTA = RGBColor(0x12, 0x30, 0x5B)


def montar(documento):
    """Margens e fonte de documento que vai ser impresso e assinado."""
    secao = documento.sections[0]
    secao.top_margin = Cm(2.5)
    secao.bottom_margin = Cm(2.5)
    secao.left_margin = Cm(2.5)
    secao.right_margin = Cm(2.5)

    normal = documento.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15


def escrever_com_negrito(paragrafo, texto):
    """Converte **assim** em negrito de verdade, e limpa o resto do markdown."""
    texto = re.sub(r'`([^`]*)`', r'\1', texto)
    texto = texto.replace('<br>', '').replace('*(opcional, mas recomendado)*', '(opcional)')

    for pedaco in re.split(r'(\*\*[^*]+\*\*)', texto):
        if not pedaco:
            continue
        if pedaco.startswith('**') and pedaco.endswith('**'):
            paragrafo.add_run(pedaco[2:-2]).bold = True
        else:
            paragrafo.add_run(pedaco)


def juntar_paragrafos(linhas):
    """
    Une as linhas de um mesmo parágrafo.

    O markdown quebra a frase em 80 colunas para eu conseguir ler o texto
    e o diff. Sem juntar de volta, cada quebra dessas vira um parágrafo no
    documento e o contrato sai picotado no meio das frases.

    Título, lista, tabela, citação e linha de assinatura ficam como estão:
    neles a quebra é o que separa um item do outro.
    """
    sozinhas = ('#', '- ', '* ', '|', '> ', '---', '===')
    saida = []
    acumulado = []

    def descarregar():
        if acumulado:
            saida.append(' '.join(acumulado))
            acumulado.clear()

    for linha in linhas:
        texto = linha.strip()

        if not texto:
            descarregar()
            continue

        if texto.startswith(sozinhas) or '____' in texto or texto.startswith('<br'):
            descarregar()
            saida.append(texto)
            continue

        acumulado.append(texto)

    descarregar()
    return saida


def converter(linhas, documento):
    linhas = juntar_paragrafos(linhas)
    dentro_de_tabela = False
    linhas_da_tabela = []

    def fechar_tabela():
        nonlocal dentro_de_tabela, linhas_da_tabela
        if not linhas_da_tabela:
            dentro_de_tabela = False
            return

        colunas = len(linhas_da_tabela[0])
        tabela = documento.add_table(rows=0, cols=colunas)
        tabela.style = 'Light Grid Accent 1'

        for indice, celulas in enumerate(linhas_da_tabela):
            linha = tabela.add_row().cells
            for coluna, conteudo in enumerate(celulas[:colunas]):
                linha[coluna].text = ''
                p = linha[coluna].paragraphs[0]
                escrever_com_negrito(p, conteudo)
                if indice == 0:
                    for run in p.runs:
                        run.bold = True

        documento.add_paragraph()
        dentro_de_tabela = False
        linhas_da_tabela = []

    for linha in linhas:
        crua = linha.rstrip('\n')
        texto = crua.strip()

        if texto.startswith('|'):
            celulas = [c.strip() for c in texto.strip('|').split('|')]
            # A linha de traços que separa o cabeçalho não vira conteúdo.
            if all(set(c) <= set('-: ') for c in celulas):
                continue
            dentro_de_tabela = True
            linhas_da_tabela.append(celulas)
            continue

        if dentro_de_tabela:
            fechar_tabela()

        if not texto:
            continue

        if texto.startswith('## '):
            documento.add_page_break() if documento.paragraphs else None
            p = documento.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(texto[3:].strip())
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = TINTA
            p.paragraph_format.space_after = Pt(16)
            continue

        if texto.startswith('### '):
            p = documento.add_paragraph()
            run = p.add_run(texto[4:].strip())
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = TINTA
            p.paragraph_format.space_before = Pt(14)
            continue

        if texto.startswith('> '):
            p = documento.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            escrever_com_negrito(p, texto[2:])
            for run in p.runs:
                run.italic = True
            continue

        if texto.startswith('- ') or texto.startswith('* '):
            p = documento.add_paragraph(style='List Bullet')
            escrever_com_negrito(p, texto[2:])
            continue

        if texto.startswith('---') or texto.startswith('==='):
            continue

        p = documento.add_paragraph()
        escrever_com_negrito(p, texto)
        # Justificado, menos nas linhas de assinatura.
        if '____' not in texto:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if dentro_de_tabela:
        fechar_tabela()


def main():
    if not os.path.exists(ORIGEM):
        print(f'não achei {ORIGEM}')
        print('gere primeiro a versão preenchível a partir de docs/contrato-modelo.md')
        return 1

    with open(ORIGEM, encoding='utf-8') as arquivo:
        linhas = arquivo.readlines()

    documento = Document()
    montar(documento)
    converter(linhas, documento)
    documento.save(DESTINO)

    faltando = sum(open(ORIGEM, encoding='utf-8').read().count(marca)
                   for marca in ('[___]', '[nome', '[endereço', '[estado', '[profissão', '[nacionalidade'))

    print(f'gerado: {DESTINO}')
    if faltando:
        print(f'\nATENÇÃO: {faltando} lacunas ainda por preencher no contrato.')
        print('Procure por colchetes no arquivo antes de mandar para ela.')
    return 0


if __name__ == '__main__':
    sys.exit(main())
